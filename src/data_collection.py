"""
data_collection.py — Resume and Job Description parsing module.
Handles text extraction from PDF, DOCX, and plain text sources.
"""

from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

try:
    import pypdf  # type: ignore[import]
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    logger.warning("pypdf not installed — PDF parsing unavailable")

try:
    import docx  # type: ignore[import]
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed — DOCX parsing unavailable")


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class Resume:
    """Represents a parsed resume document."""

    raw_text: str
    filename: str
    entities: dict[str, list[str]] = field(default_factory=dict)
    clean_text: str = ""
    skills: list[str] = field(default_factory=list)
    education: list[str] = field(default_factory=list)
    experience_years: float = 0.0
    metadata: dict[str, Any] = field(default_factory=dict)

    def __post_init__(self) -> None:
        if not self.clean_text:
            self.clean_text = ResumeParser.clean_text(self.raw_text)


@dataclass
class JobDescription:
    """Represents a parsed job description."""

    raw_text: str
    title: str
    required_skills: list[str] = field(default_factory=list)
    preferred_skills: list[str] = field(default_factory=list)
    min_experience: float = 0.0
    education_level: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Resume parser
# ---------------------------------------------------------------------------


class ResumeParser:
    """Parse resume documents from multiple file formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

    def parse_file(self, file_path: str | Path) -> Resume:
        """Parse a resume from a file on disk.

        Args:
            file_path: Path to the resume file.

        Returns:
            Parsed Resume object.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file extension is unsupported.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file extension: {ext}")

        raw_text = self._extract_text(path)
        logger.info("Parsed resume from %s (%d chars)", path.name, len(raw_text))
        return self.parse_text(raw_text, filename=path.name)

    def parse_text(self, text: str, filename: str = "unknown.txt") -> Resume:
        """Parse a resume from raw text.

        Args:
            text: Raw resume text.
            filename: Source filename for metadata.

        Returns:
            Parsed Resume object.
        """
        clean = self.clean_text(text)
        return Resume(
            raw_text=text,
            filename=filename,
            clean_text=clean,
            metadata={"source": filename, "char_count": len(text)},
        )

    def _extract_text(self, path: Path) -> str:
        """Dispatch extraction based on file extension."""
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._extract_pdf(path)
        if ext in {".docx", ".doc"}:
            return self._extract_docx(path)
        # Plain text / markdown
        return path.read_text(encoding="utf-8", errors="replace")

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from a PDF file using pypdf.

        Args:
            path: Path to the PDF file.

        Returns:
            Extracted text.

        Raises:
            ImportError: If pypdf is not installed.
        """
        if not HAS_PYPDF:
            raise ImportError("pypdf is required to parse PDF files: pip install pypdf")

        pages: list[str] = []
        with pypdf.PdfReader(str(path)) as reader:
            for page in reader.pages:
                page_text = page.extract_text() or ""
                pages.append(page_text)
        return "\n".join(pages)

    def _extract_docx(self, path: Path) -> str:
        """Extract text from a DOCX file using python-docx.

        Args:
            path: Path to the DOCX file.

        Returns:
            Extracted text.

        Raises:
            ImportError: If python-docx is not installed.
        """
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is required to parse DOCX files: pip install python-docx"
            )

        document = docx.Document(str(path))
        paragraphs = [para.text for para in document.paragraphs]
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(p for p in paragraphs if p.strip())

    @staticmethod
    def clean_text(text: str) -> str:
        """Normalize and sanitize resume text.

        Steps:
        - Remove non-printable / control characters
        - Normalize unicode characters
        - Normalize bullet point characters to a dash
        - Collapse excess whitespace
        - Mask email addresses and phone numbers

        Args:
            text: Raw input text.

        Returns:
            Cleaned text string.
        """
        if not text:
            return ""

        # Normalize unicode
        text = unicodedata.normalize("NFKC", text)

        # Remove non-printable control characters (except newline/tab)
        text = "".join(
            ch if (unicodedata.category(ch) != "Cc" or ch in "\n\t") else " "
            for ch in text
        )

        # Normalize common bullet characters to ASCII dash
        bullet_chars = "•·‣▸▹◦➢➤►▶→»"
        for ch in bullet_chars:
            text = text.replace(ch, "-")

        # Mask email addresses
        text = re.sub(
            r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
            "[EMAIL]",
            text,
        )

        # Mask phone numbers (various formats)
        text = re.sub(
            r"(\+?1[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}",
            "[PHONE]",
            text,
        )

        # Normalize whitespace: collapse multiple spaces/tabs; preserve newlines
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
        text = "\n".join(lines)

        # Collapse multiple blank lines into a single blank line
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()


# ---------------------------------------------------------------------------
# Job description parser
# ---------------------------------------------------------------------------


class JobDescriptionParser:
    """Parse job descriptions from text or file."""

    # Patterns to identify required vs preferred sections
    _REQUIRED_SECTION = re.compile(
        r"(required|must.have|minimum qualifications?)[:\s]*", re.IGNORECASE
    )
    _PREFERRED_SECTION = re.compile(
        r"(preferred|nice.to.have|bonus|desired)[:\s]*", re.IGNORECASE
    )
    _EXPERIENCE_PATTERN = re.compile(
        r"(\d+)\+?\s*(?:to\s*\d+\s*)?years?\s+(?:of\s+)?(?:relevant\s+)?experience",
        re.IGNORECASE,
    )
    _EDU_LEVELS = {
        "phd": "PhD",
        "doctorate": "PhD",
        "master": "Master's",
        "msc": "Master's",
        "bachelor": "Bachelor's",
        "bsc": "Bachelor's",
        "associate": "Associate's",
        "high school": "High School",
    }

    def parse_text(self, text: str, title: str = "Unknown Position") -> JobDescription:
        """Parse a job description from raw text.

        Args:
            text: Raw job description text.
            title: Job title (defaults to "Unknown Position").

        Returns:
            Parsed JobDescription object.
        """
        min_exp = self._extract_min_experience(text)
        edu_level = self._extract_education_level(text)

        return JobDescription(
            raw_text=text,
            title=title,
            min_experience=min_exp,
            education_level=edu_level,
            metadata={"char_count": len(text)},
        )

    def parse_file(self, file_path: str | Path) -> JobDescription:
        """Parse a job description from a plain-text or markdown file.

        Args:
            file_path: Path to the job description file.

        Returns:
            Parsed JobDescription object.

        Raises:
            FileNotFoundError: If the file does not exist.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Job description file not found: {path}")

        text = path.read_text(encoding="utf-8", errors="replace")
        title = path.stem.replace("_", " ").replace("-", " ").title()
        logger.info("Parsed job description from %s", path.name)
        return self.parse_text(text, title=title)

    def _extract_min_experience(self, text: str) -> float:
        """Extract minimum years of experience from text."""
        matches = self._EXPERIENCE_PATTERN.findall(text)
        if matches:
            return float(min(int(m) for m in matches))
        return 0.0

    def _extract_education_level(self, text: str) -> str:
        """Detect the minimum education level mentioned in text."""
        lower = text.lower()
        for keyword, level in self._EDU_LEVELS.items():
            if keyword in lower:
                return level
        return ""
